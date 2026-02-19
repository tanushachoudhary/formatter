import base64
import os
import tempfile

from docx import Document
from docx.shared import Inches

from utils.docx_to_images import docx_to_page_images, docx_to_page_images_base64, ocr_page_images
from utils.formatter import (
    clear_document_body,
    force_legal_run_format_document,
    force_single_column,
    inject_blocks,
    remove_trailing_empty_and_noise,
)
from utils.llm_formatter import format_text_with_llm
from utils.style_extractor import (
    _paragraph_has_bottom_border,
    extract_document_blueprint,
    extract_styles,
    load_extracted_styles,
    save_document_blueprint,
    save_extracted_styles,
)

# Summons-style page margins (generous, like formal legal documents)
DEFAULT_TOP_MARGIN_IN = 1.25
DEFAULT_BOTTOM_MARGIN_IN = 1.25
DEFAULT_LEFT_MARGIN_IN = 1.25
DEFAULT_RIGHT_MARGIN_IN = 1.25



def _project_dir():
    return os.path.dirname(os.path.abspath(__file__))


def get_document_preview_text(docx_path: str) -> str:
    """Build a plain-text preview of the formatted DOCX for display before download.
    Paragraphs with only a bottom border (section underlines) are emitted as [SECTION_UNDERLINE]."""
    doc = Document(docx_path)
    lines = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text and _paragraph_has_bottom_border(para):
            lines.append("[SECTION_UNDERLINE]")
        else:
            lines.append(text if text else "")
    return "\n\n".join(lines).strip()


def extract_and_store_styles(template_file) -> dict:
    """Extract styles from the uploaded DOCX and save to JSON. Returns the style schema."""
    doc = Document(template_file)
    schema = extract_styles(doc)
    save_extracted_styles(schema, base_dir=_project_dir())
    blueprint = extract_document_blueprint(doc)
    save_document_blueprint(blueprint, base_dir=_project_dir())
    return schema


def process_document(generated_text, template_file):
    """
    Input 1: Uploaded DOCX template (desired styles and formatting).
    Input 2: Raw legal text (unformatted).
    Segment and render entire text using template styles (no slot-fill).
    Template is also converted to page images and sent to the LLM when possible (vision).
    """
    project_dir = _project_dir()
    doc = Document(template_file)

    schema = extract_styles(doc)
    save_extracted_styles(schema, base_dir=project_dir)

    # Convert document to images (each page → image), then send to LLM for formatting reference.
    # Template may have multi-column layout; convert a single-column copy so each page image is one column (not 3 side-by-side).
    template_path = None
    single_column_path = None
    template_page_images = []
    template_page_ocr_texts = []
    try:
        template_file.seek(0)
        data = template_file.read()
        template_file.seek(0)
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(data)
            tmp.flush()
            template_path = tmp.name
        doc_for_images = Document(template_path)
        force_single_column(doc_for_images)
        fd, single_column_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        doc_for_images.save(single_column_path)
        page_bytes = docx_to_page_images(single_column_path, dpi=150, max_pages=15)
        if page_bytes:
            template_page_images = [base64.b64encode(b).decode("ascii") for b in page_bytes]
            schema["template_page_images"] = template_page_images
            template_page_ocr_texts = ocr_page_images(page_bytes)
            if template_page_ocr_texts and any(t.strip() for t in template_page_ocr_texts):
                schema["template_page_ocr_texts"] = template_page_ocr_texts
    except Exception:
        pass
    for path in (single_column_path, template_path):
        if path and os.path.isfile(path):
            try:
                os.unlink(path)
            except OSError:
                pass

    blocks = format_text_with_llm(
        generated_text,
        schema,
        use_slot_fill=False,
        template_page_images=template_page_images,
        template_page_ocr_texts=template_page_ocr_texts if template_page_ocr_texts else None,
    )

    clear_document_body(doc)
    force_single_column(doc)
    inject_blocks(
        doc,
        blocks,
        style_map=schema["style_map"],
        style_formatting=schema.get("style_formatting", {}),
        line_samples=schema.get("line_samples", []),
        section_heading_samples=schema.get("section_heading_samples", []),
        template_structure=None,
        numbered_num_id=schema.get("numbered_num_id"),
        numbered_ilvl=schema.get("numbered_ilvl", 0),
    )
    force_legal_run_format_document(doc)
    remove_trailing_empty_and_noise(doc)

    output_path = os.path.join(project_dir, "output", "formatted_output.docx")
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    preview_text = get_document_preview_text(output_path)
    return output_path, preview_text
