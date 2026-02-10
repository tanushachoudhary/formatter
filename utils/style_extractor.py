"""Extract styles and formatting from a DOCX and store as JSON."""

import json
import os

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.shared import Length
from docx.table import Table
from docx.text.paragraph import Paragraph

STORE_DIR = "output"
EXTRACTED_STYLES_FILE = "extracted_styles.json"
EXTRACTED_STYLE_GUIDE_FILE = "extracted_style_guide.txt"
EXTRACTED_BLUEPRINT_FILE = "document_blueprint.json"

PREFERRED_HEADING_1 = ("Heading 1", "Title", "Titre 1")
PREFERRED_HEADING_2 = ("Heading 2", "Subtitle", "Titre 2", "Section")
PREFERRED_NORMAL = ("Normal", "Body Text", "Paragraphe")
PREFERRED_LIST = ("List Number", "List Paragraph", "List")


def _get_paragraph_style_names(doc):
    return [
        s.name for s in doc.styles
        if s.type == WD_STYLE_TYPE.PARAGRAPH
    ]


def clone_styles(src_doc: Document, dst_doc: Document) -> None:
    """Clone paragraph style definitions from template (src) to destination (dst).
    Use when building a new document that must match template layout — then assign
    paragraph.style = template_style_name so Word handles layout, numbering, spacing.
    Does not copy numbering definitions; for that see numbering_part (Upgrade 2)."""
    dst_names = {s.name for s in dst_doc.styles if s.type == WD_STYLE_TYPE.PARAGRAPH}
    for style in src_doc.styles:
        if style.type != WD_STYLE_TYPE.PARAGRAPH:
            continue
        if style.name in dst_names:
            continue
        try:
            new_style = dst_doc.styles.add_style(style.name, WD_STYLE_TYPE.PARAGRAPH)
            if getattr(style, "base_style", None) and style.base_style is not None:
                base_name = style.base_style.name if hasattr(style.base_style, "name") else str(style.base_style)
                if base_name in dst_names:
                    new_style.base_style = dst_doc.styles[base_name]
        except Exception:
            pass


def _pick_style(available, preferred_names, fallback_names=None):
    for name in preferred_names:
        if name in available:
            return name
    if fallback_names:
        for name in fallback_names:
            if name in available:
                return name
    return available[0] if available else None


def _length_pt(val):
    """Serialize a Length or None to points (float) for JSON."""
    if val is None:
        return None
    if isinstance(val, Length):
        return getattr(val, "pt", None)
    return None


def _enum_name(val):
    """Get enum member name for JSON (e.g. WD_UNDERLINE.DOTTED -> 'DOTTED')."""
    if val is None:
        return None
    if hasattr(val, "name"):
        return val.name
    return str(val)


def _extract_paragraph_format(pf):
    """Extract paragraph format to a JSON-serializable dict (Word document features: alignment, spacing, indents, line_spacing, keep_*, page_break_before)."""
    if pf is None:
        return {}
    out = {}
    for attr in ("alignment", "space_before", "space_after", "left_indent", "right_indent", "first_line_indent", "line_spacing"):
        try:
            val = getattr(pf, attr, None)
            if val is None:
                continue
            if attr == "alignment":
                out[attr] = _enum_name(val)
            elif attr == "line_spacing":
                if isinstance(val, Length):
                    out[attr] = val.pt
                elif isinstance(val, (int, float)):
                    out[attr] = val
                else:
                    out[attr] = None
            else:
                out[attr] = _length_pt(val) if hasattr(val, "pt") else val
        except Exception:
            pass
    try:
        rule = getattr(pf, "line_spacing_rule", None)
        if rule is not None:
            out["line_spacing_rule"] = _enum_name(rule)
    except Exception:
        pass
    for attr in ("page_break_before", "keep_with_next", "keep_together"):
        try:
            val = getattr(pf, attr, None)
            if val is not None:
                out[attr] = bool(val)
        except Exception:
            pass
    try:
        if getattr(pf, "widow_control", None) is not None:
            out["widow_control"] = bool(pf.widow_control)
    except Exception:
        pass
    try:
        tab_stops = getattr(pf, "tab_stops", None)
        if tab_stops is not None and hasattr(tab_stops, "__iter__"):
            stops = []
            for ts in tab_stops:
                try:
                    pos = _length_pt(getattr(ts, "position", None))
                    align = _enum_name(getattr(ts, "alignment", None))
                    leader = _enum_name(getattr(ts, "leader", None))
                    if pos is not None or align or leader:
                        stops.append({"position_pt": pos, "alignment": align, "leader": leader})
                except Exception:
                    pass
            if stops:
                out["tab_stops"] = stops
    except Exception:
        pass
    return out


def _extract_run_format(run):
    """Extract run/font format to a JSON-serializable dict (bold, italic, underline, font name/size, color)."""
    if run is None:
        return {}
    return _extract_run_format_from_font(run.font)


def _format_from_style_definition(doc: Document, style_name: str) -> dict | None:
    """Extract paragraph_format (and optionally font) from the style definition when no paragraph uses it."""
    try:
        style = doc.styles[style_name]
    except KeyError:
        return None
    if style.type != WD_STYLE_TYPE.PARAGRAPH:
        return None
    pf = {}
    run_fmt = {}
    try:
        pf = _extract_paragraph_format(style.paragraph_format)
    except Exception:
        pass
    try:
        # Style's font is on the style element (rPr); get first run property set
        if hasattr(style, "font") and style.font:
            run_fmt = _extract_run_format_from_font(style.font)
    except Exception:
        pass
    if not pf and not run_fmt:
        return None
    return {"paragraph_format": pf, "run_format": run_fmt}


def _extract_run_format_from_font(font) -> dict:
    """Extract run format from a Font object (e.g. from a style): bold, italic, underline, font name/size, color."""
    if font is None:
        return {}
    out = {}
    try:
        if font.bold is not None:
            out["bold"] = font.bold
    except Exception:
        pass
    try:
        if font.italic is not None:
            out["italic"] = font.italic
    except Exception:
        pass
    try:
        u = font.underline
        if u is not None:
            out["underline"] = _enum_name(u) if hasattr(u, "name") else (True if u else False)
    except Exception:
        pass
    try:
        if font.name is not None:
            out["name"] = font.name
    except Exception:
        pass
    try:
        if font.size is not None and hasattr(font.size, "pt"):
            out["size_pt"] = font.size.pt
    except Exception:
        pass
    try:
        if hasattr(font, "color") and font.color and getattr(font.color, "rgb", None):
            rgb = font.color.rgb
            if rgb is not None and len(rgb) >= 3:
                out["color_rgb_hex"] = "%02x%02x%02x" % (int(rgb[0]), int(rgb[1]), int(rgb[2]))
    except Exception:
        pass
    return out


def _merge_format(base: dict, override: dict) -> dict:
    """Merge override into base; override wins when both have a key. For nested dicts, merge recursively."""
    out = dict(base)
    for k, v in override.items():
        if k not in out or v is not None and v != "" and (not isinstance(v, dict) or v):
            if isinstance(v, dict) and isinstance(out.get(k), dict):
                out[k] = _merge_format(out[k], v)
            else:
                out[k] = v
    return out


# Default spacing/indent when template has none. Use 0 so output matches tight,
# single-spaced templates; template-defined spacing is always preserved.
DEFAULT_SPACE_AFTER_PT = 0.0
DEFAULT_SPACE_BEFORE_HEADING_PT = 0.0
DEFAULT_FIRST_LINE_INDENT_PT = 0.0
DEFAULT_NUMBERED_LEFT_INDENT_PT = 0.0
DEFAULT_NUMBERED_FIRST_LINE_INDENT_PT = 0.0


def _sample_formatting_per_style(doc: Document) -> dict:
    """For each paragraph style used in the doc (including in table cells), sample one paragraph and extract its formatting."""
    style_formatting = {}
    seen_styles = set()
    for para, _tid, _r, _c in iter_body_blocks(doc):
        try:
            style_name = para.style.name if para.style else None
        except Exception:
            style_name = None
        if not style_name or style_name in seen_styles:
            continue
        seen_styles.add(style_name)
        pf = _extract_paragraph_format(para.paragraph_format)
        run_fmt = {}
        if para.runs:
            run_fmt = _extract_run_format(para.runs[0])
        if pf or run_fmt:
            style_formatting[style_name] = {
                "paragraph_format": pf,
                "run_format": run_fmt,
            }
    return style_formatting


def _enrich_style_formatting_from_definitions(doc: Document, style_map: dict, style_formatting: dict) -> dict:
    """Fill in formatting from style definitions. Preserve template spacing and indentation:
    use paragraph sample when present, otherwise use style definition so output matches the template."""
    result = dict(style_formatting)
    for block_type, style_name in style_map.items():
        if not style_name:
            continue
        from_para = result.get(style_name, {})
        from_def = _format_from_style_definition(doc, style_name)
        if from_def:
            merged = _merge_format(from_def, from_para)
            # Keep merged paragraph_format as-is so spacing/indent from template (style or paragraph) is preserved
            result[style_name] = merged
        elif style_name not in result:
            result[style_name] = {"paragraph_format": {}, "run_format": {}}
    return result


def _apply_default_spacing_and_indent(style_map: dict, style_formatting: dict) -> dict:
    """Ensure paragraph and numbered styles have sensible spacing and indent when missing."""
    result = dict(style_formatting)
    normal_style = style_map.get("paragraph")
    list_style = style_map.get("numbered")
    if normal_style and normal_style in result:
        pf = result[normal_style].get("paragraph_format") or {}
        pf = dict(pf)
        if pf.get("space_after") is None:
            pf["space_after"] = DEFAULT_SPACE_AFTER_PT
        if pf.get("first_line_indent") is None:
            pf["first_line_indent"] = DEFAULT_FIRST_LINE_INDENT_PT
        result[normal_style] = {**result[normal_style], "paragraph_format": pf}
    if list_style and list_style in result:
        pf = result[list_style].get("paragraph_format") or {}
        pf = dict(pf)
        if pf.get("left_indent") is None:
            pf["left_indent"] = DEFAULT_NUMBERED_LEFT_INDENT_PT
        if pf.get("first_line_indent") is None:
            pf["first_line_indent"] = DEFAULT_NUMBERED_FIRST_LINE_INDENT_PT
        if pf.get("space_after") is None:
            pf["space_after"] = DEFAULT_SPACE_AFTER_PT
        result[list_style] = {**result[list_style], "paragraph_format": pf}
    return result


def _format_spec_to_lines(style_name: str, fmt: dict) -> list[str]:
    """Turn one style's paragraph_format + run_format into plain-text bullet lines."""
    lines = [f"Style: {style_name}", ""]
    pf = fmt.get("paragraph_format") or {}
    rf = fmt.get("run_format") or {}
    # Run/font
    run_parts = []
    if rf.get("bold"):
        run_parts.append("bold")
    if rf.get("italic"):
        run_parts.append("italic")
    if rf.get("underline"):
        u = rf["underline"]
        run_parts.append(f"underline ({u})" if isinstance(u, str) else "underline")
    if rf.get("name"):
        run_parts.append(f"font: {rf['name']}")
    if rf.get("size_pt") is not None:
        run_parts.append(f"{rf['size_pt']}pt")
    if run_parts:
        lines.append("- " + ", ".join(run_parts))
    # Paragraph
    para_parts = []
    if pf.get("alignment"):
        para_parts.append(f"align: {pf['alignment'].lower()}")
    if pf.get("space_before") is not None:
        para_parts.append(f"space before: {pf['space_before']}pt")
    if pf.get("space_after") is not None:
        para_parts.append(f"space after: {pf['space_after']}pt")
    if pf.get("left_indent") is not None:
        para_parts.append(f"left indent: {pf['left_indent']}pt")
    if pf.get("first_line_indent") is not None:
        para_parts.append(f"first line indent: {pf['first_line_indent']}pt")
    if pf.get("line_spacing") is not None:
        para_parts.append(f"line spacing: {pf['line_spacing']}pt")
    if para_parts:
        lines.append("- " + ", ".join(para_parts))
    if not run_parts and not para_parts:
        lines.append("- (default paragraph)")
    lines.append("")
    return lines


def _is_line_paragraph(para) -> bool:
    """True if paragraph text looks like a separator line (dots, dashes, underscores, optional X)."""
    if not para or not para.text:
        return False
    t = para.text.strip()
    if not t or len(t) < 3:
        return False
    # Allow letters only for trailing X or similar; rest should be symbols/spaces
    allowed = set(" ._-= \t\u00A0")
    has_symbol = False
    for c in t:
        if c in allowed or c in ".,":
            if c not in " \t":
                has_symbol = True
        elif c.isalpha() and c.upper() == "X" and t.rstrip().endswith("X"):
            continue
        elif not c.isalnum():
            has_symbol = True
        else:
            return False
    return has_symbol


def _paragraph_has_bottom_border(para) -> bool:
    """True if paragraph has a bottom border (used for section underlines under headings)."""
    try:
        p = getattr(para, "_p", None)
        if p is None:
            return False
        pPr = p.find(qn("w:pPr"))
        if pPr is None:
            return False
        pBdr = pPr.find(qn("w:pBdr"))
        if pBdr is None:
            return False
        return pBdr.find(qn("w:bottom")) is not None
    except Exception:
        return False


def _is_signature_line_paragraph(para) -> bool:
    """True if paragraph looks like a signature underline (mostly underscores)."""
    if not para or not para.text:
        return False
    t = para.text.strip()
    if not t or len(t) < 2:
        return False
    without_underscore = t.replace("_", "").replace(" ", "")
    return len(without_underscore) == 0 or (len(without_underscore) <= 2 and "_" in t)


def _extract_line_samples(doc: Document) -> list[dict]:
    """Extract paragraphs that look like separator/signature lines from the template."""
    samples = []
    seen_text = set()
    for para in doc.paragraphs:
        if not _is_line_paragraph(para):
            continue
        text = para.text.strip()
        if not text or text in seen_text:
            continue
        seen_text.add(text)
        alignment = None
        try:
            pf = para.paragraph_format
            if pf and pf.alignment is not None:
                alignment = _enum_name(pf.alignment)
        except Exception:
            pass
        samples.append({"text": text, "alignment": alignment})
    return samples


def build_style_guide(style_map: dict, style_formatting: dict, all_style_names: list = None) -> str:
    """Build plain-text listing of template styles and their formatting (font, bold, alignment, indent, etc.)."""
    lines = [
        "Extracted styles and formatting (from uploaded DOCX template)",
        "",
        "Goal: Apply these styles to raw text so the output has the same structure and formatting as the template. Match how the template uses each style (title, section headings, body, numbered lists, etc.). Use the exact style name as block_type.",
        "",
        "Application rules (match template structure)",
        "",
    ]
    h1 = style_map.get("heading")
    h2 = style_map.get("section_header")
    normal = style_map.get("paragraph")
    list_style = style_map.get("numbered")
    if h1:
        lines.append(f"- Document title / main heading (court name, case caption, or document title like MOTION TO RESTORE) -> use style: {h1}")
    if h2:
        lines.append(f"- Section headers / subheadings (e.g. VERIFIED COMPLAINT, MEMORANDUM, FACTS, ARGUMENT, CONCLUSION, RELIEF REQUESTED, party captions) -> use style: {h2}")
    if normal:
        lines.append(f"- Body paragraphs (narrative, legal argument, facts, verification text, general content) -> use style: {normal}")
    if list_style:
        lines.append(f"- Numbered or bulleted list items (allegations, motion grounds, relief items, any list where the template uses this style) -> use style: {list_style} (numbers 1., 2., 3. added automatically when applicable).")
    names_for_caption = all_style_names or list(style_formatting.keys()) if style_formatting else []
    caption_like = [n for n in names_for_caption if n and ("caption" in n.lower() or "right" in n.lower() or "index" in n.lower())]
    if caption_like:
        lines.append(f"- Case number / date / right-aligned caption -> use style {caption_like[0]} where the template uses it.")
    lines.append("- Addresses (TO:, court or party address lines): use the same style as in the template (often Normal; one block per line if the template breaks them).")
    lines.append("- Signature block (attorney name, firm, address, phone): use the template style for that block; use block_type signature_line for the underline line.")
    lines.append("- Separator lines (dashes/dots ending in X) -> use block_type line and put the line characters in text.")
    lines.append("- Signature underlines -> use block_type signature_line (optional label in text).")
    lines.append("- Section underlines (solid line under a cause of action or heading) -> use block_type section_underline with empty text.")
    lines.append("- Page break -> use block_type page_break with empty text before each new major section (as in the template).")
    lines.append("")
    lines.append("Style definitions (font, alignment, indent, etc.)")
    lines.append("")
    names = list(all_style_names) if all_style_names else list(style_formatting.keys()) if style_formatting else []
    if not names and style_map:
        names = list(set(style_map.values()))
    for style_name in names:
        fmt = style_formatting.get(style_name, {})
        lines.extend(_format_spec_to_lines(style_name, fmt))
    if not names:
        lines.append("No styles in template.")
    return "\n".join(lines).strip()


def _section_label(text: str, index: int) -> str:
    """Return a short section label for a template paragraph (for dynamic prompts)."""
    t = (text or "").strip().lower()
    if not t:
        return "continuation"
    if "supreme court" in t or "county of" in t[:30]:
        return "court_header"
    if "plaintiff" in t and ("against" in t or len(t) < 30):
        return "case_caption_parties"
    if "against" in t and len(t) < 20:
        return "case_caption_parties"
    if "defendant" in t and len(t) < 50:
        return "case_caption_parties"
    if "index no" in t or "index no." in t or "docket no" in t:
        return "case_caption_index"
    if "notice of motion" in t or "to restore" in t or "affirmation in support" in t or "affidavit of service" in t or "summons" in t or "complaint" in t:
        if len(t) < 80:
            return "document_title"
    if "counselors" in t or "c o u n s e l o r s" in t:
        return "counselors_header"
    if "please take notice" in t or "take further notice" in t:
        return "body_notice"
    if t.startswith("dated") or t.startswith("dated:"):
        return "date_line"
    if "________" in t or ("_" in t and len(t.strip()) < 5):
        return "signature_line"
    if "attorneys for" in t and ("plaintiff" in t or "defendant" in t):
        return "signature_block"
    if t.strip() == "to:" or (t.startswith("to:") and len(t) < 10):
        return "recipients_header"
    if "state of new york" in t and ")" in t and "ss." in t:
        return "affidavit_opening"
    if "being duly sworn" in t or "duly sworn" in t:
        return "affidavit_body"
    return "body"


def build_section_formatting_prompts(template_content: list, style_formatting: dict) -> str:
    """Build per-section formatting instructions for the LLM (dynamic prompt per section type)."""
    if not template_content or not style_formatting:
        return ""
    seen_sections = {}
    lines = [
        "Per-section formatting (apply the following to the corresponding parts of the raw text):",
        "",
    ]
    for i, item in enumerate(template_content):
        style_name = item.get("style") or "Normal"
        text = (item.get("text") or "").strip()
        section = _section_label(text, i)
        fmt = style_formatting.get(style_name, {})
        pf = fmt.get("paragraph_format", {})
        run_fmt = fmt.get("run_format", {})
        alignment = pf.get("alignment", "")
        bold = run_fmt.get("bold", False)
        parts = [f"Style: {style_name}"]
        if alignment:
            parts.append(f"alignment: {alignment.lower()}")
        if bold:
            parts.append("bold")
        fmt_desc = ", ".join(parts)
        # One prompt line per section type (first occurrence) or per distinct style in that section
        key = (section, style_name)
        if key not in seen_sections:
            seen_sections[key] = True
            snippet = (text[:50] + "…") if len(text) > 50 else text
            if not snippet:
                snippet = "(empty line)"
            lines.append(f"• {section}: Use {fmt_desc}. Example content: \"{snippet}\"")
    if len(lines) <= 2:
        return ""
    return "\n".join(lines)


def iter_body_blocks(doc: Document):
    """
    Yield (paragraph, table_id, row, col) in document order.
    table_id is None for top-level body paragraphs; inside tables, table_id is an int (0, 1, ...),
    and row/col are 0-based cell indices. This includes paragraphs inside table cells so caption
    content is not missed when the template uses a table for the caption.
    """
    body = doc.element.body
    qp, qt = qn("w:p"), qn("w:tbl")
    table_counter = 0
    for child in body.iterchildren():
        if child.tag == qp:
            yield Paragraph(child, doc), None, None, None
        elif child.tag == qt:
            tbl = Table(child, doc)
            for ri, row in enumerate(tbl.rows):
                for ci, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        yield para, table_counter, ri, ci
            table_counter += 1
        # skip sectPr and any other elements


def extract_tables(doc: Document) -> list[dict]:
    """
    Extract table structure from the document: for each table, record index, dimensions (rows × cols),
    and a short preview of cell contents. Tables are in document order (same order as in iter_body_blocks).
    """
    body = doc.element.body
    qp, qt = qn("w:p"), qn("w:tbl")
    tables_out = []
    table_index = 0
    for child in body.iterchildren():
        if child.tag == qt:
            tbl = Table(child, doc)
            num_rows = len(tbl.rows)
            num_cols = len(tbl.columns) if tbl.rows else 0
            cell_preview = []
            for row in tbl.rows:
                row_preview = []
                for cell in row.cells:
                    parts = []
                    for para in cell.paragraphs:
                        t = (para.text or "").strip()
                        if t:
                            parts.append(t[:60] + ("…" if len(t) > 60 else ""))
                    row_preview.append(" | ".join(parts) if parts else "")
                cell_preview.append(row_preview)
            tables_out.append({
                "table_index": table_index,
                "rows": num_rows,
                "cols": num_cols,
                "cell_preview": cell_preview,
            })
            table_index += 1
    return tables_out


def get_template_content_with_styles(doc: Document, max_paragraphs: int = 120, max_text_per_para: int = 400) -> list[dict]:
    """Extract each paragraph's style name and text so the LLM can see how the template was formatted."""
    out = []
    for i, (para, _tid, _r, _c) in enumerate(iter_body_blocks(doc)):
        if i >= max_paragraphs:
            break
        try:
            style_name = para.style.name if para.style else None
        except Exception:
            style_name = None
        text = (para.text or "").strip()
        if max_text_per_para and len(text) > max_text_per_para:
            text = text[: max_text_per_para] + "..."
        out.append({"style": style_name or "Normal", "text": text})
    return out


def _extract_section_heading_samples(doc: Document) -> list[str]:
    """Extract normalized text of paragraphs that have 'page break before' in the template (for section breaks)."""
    seen = set()
    out = []
    for para, _tid, _r, _c in iter_body_blocks(doc):
        try:
            if not getattr(para.paragraph_format, "page_break_before", None):
                continue
        except Exception:
            continue
        text = (para.text or "").strip()
        if not text or len(text) < 2:
            continue
        key = text.lower().strip()
        if key not in seen:
            seen.add(key)
            out.append(key)
    return out


def _infer_section_type(hint: str, block_kind: str) -> str:
    """Infer legal document section from block hint (for layout enforcement)."""
    t = (hint or "").strip().lower()
    if block_kind in ("line", "signature_line"):
        return "separator"
    if not t or t == "(empty)":
        return "body"
    if any(x in t for x in ("supreme court", "county of", "index no", "index number", "notice of motion", "to restore", "affirmation in support", "affidavit of service", "c o u n s e l o r s", "counselors:", "plaintiff", "defendant", "-against-", "against")):
        return "caption"
    if any(x in t for x in ("please take notice", "take further notice", "for an order")):
        return "motion_notice"
    if any(x in t for x in ("attorneys for", "attorney for", "law firm", "esq.", "pllc", "p.c.")) and any(c.isdigit() for c in t):
        return "attorney_signature"
    if t.startswith("to:") or (len(t) < 5 and "to" in t):
        return "to_section"
    if any(x in t for x in ("affirms the following", "respectfully submitted", "it is respectfully", "wherefore")):
        return "affirmation"
    if any(x in t for x in ("duly sworn", "being duly sworn", "under the penalties of perjury")):
        return "affidavit"
    if any(x in t for x in ("sworn to before me", "notary public", "state of ", "county of ")) and len(t) < 120:
        return "notary"
    if "dated:" in t and len(t) < 80:
        return "body"
    return "body"


def extract_template_structure(doc: Document, max_paragraphs: int = 500) -> list[dict]:
    """
    Extract the exact structure of the template: one block spec per paragraph, in document order.
    Includes paragraphs inside table cells (e.g. caption) so slot-fill has the right number of slots.
    Each block has: style, paragraph_format, run_format, block_kind, section_type, template_text,
    page_break_before, hint; and optionally table_id, row, col when inside a table.
    """
    block_specs = []
    for i, (para, table_id, row, col) in enumerate(iter_body_blocks(doc)):
        if i >= max_paragraphs:
            break
        try:
            style_name = para.style.name if para.style else "Normal"
        except Exception:
            style_name = "Normal"
        pf = para.paragraph_format
        paragraph_format = _extract_paragraph_format(pf) if pf else {}
        run_format = {}
        if para.runs:
            run_format = _extract_run_format(para.runs[0])
        page_break_before = False
        try:
            if pf and getattr(pf, "page_break_before", None):
                page_break_before = True
        except Exception:
            pass
        text = (para.text or "").strip()
        if _is_signature_line_paragraph(para):
            block_kind = "signature_line"
            template_text = text or "_________________________"
        elif _is_line_paragraph(para):
            block_kind = "line"
            template_text = text or "----------------------------------------------------------------------X"
        elif not text and _paragraph_has_bottom_border(para):
            block_kind = "section_underline"
            template_text = None
        else:
            block_kind = "paragraph"
            template_text = None
        hint = (text[:80] + "…") if text and len(text) > 80 else (text or "(empty)")
        section_type = _infer_section_type(hint, block_kind)
        spec = {
            "style": style_name,
            "paragraph_format": paragraph_format,
            "run_format": run_format,
            "block_kind": block_kind,
            "section_type": section_type,
            "template_text": template_text,
            "page_break_before": page_break_before,
            "hint": hint,
        }
        if table_id is not None:
            spec["table_id"] = table_id
            spec["row"] = row
            spec["col"] = col
        block_specs.append(spec)
    return block_specs


def _extract_document_layout(doc: Document) -> dict:
    """Extract page layout: margins, page size, orientation, section breaks (formatting metadata only)."""
    out = {"sections": []}
    try:
        for idx, section in enumerate(doc.sections):
            sec = {}
            for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin", "page_width", "page_height"):
                try:
                    val = getattr(section, attr, None)
                    if val is not None:
                        sec[attr] = _length_pt(val) if hasattr(val, "pt") else (getattr(val, "inches", None) if val else None)
                        if sec[attr] is None and hasattr(val, "pt"):
                            sec[attr] = val.pt
                except Exception:
                    pass
            try:
                orient = getattr(section, "orientation", None)
                if orient is not None:
                    sec["orientation"] = _enum_name(orient)
            except Exception:
                pass
            if sec:
                sec["section_index"] = idx
                out["sections"].append(sec)
    except Exception:
        pass
    return out


def extract_document_blueprint(doc: Document) -> dict:
    """
    Extract the complete style and layout blueprint: styles, sections, tables, lists, document_layout.
    Formatting metadata only (no document text except as identifiers). Machine-readable schema
    for programmatic application to another document.
    """
    para_names = _get_paragraph_style_names(doc)
    style_formatting = _sample_formatting_per_style(doc)
    # Enrich from style definitions so we have full style system
    style_map = {}
    if para_names:
        heading_like = [n for n in para_names if any(kw in n.lower() for kw in ("heading", "title", "titre", "section"))]
        list_like = [n for n in para_names if "list" in n.lower() or "number" in n.lower()]
        h1 = _pick_style(para_names, PREFERRED_HEADING_1, heading_like)
        h2 = _pick_style(para_names, PREFERRED_HEADING_2, [n for n in heading_like if n != h1])
        normal = _pick_style(para_names, PREFERRED_NORMAL, para_names)
        list_style = _pick_style(para_names, PREFERRED_LIST, list_like) if list_like else normal
        style_map = {"heading": h1 or normal, "section_header": h2 or h1 or normal, "paragraph": normal, "numbered": list_style, "wherefore": h2 or h1 or normal}
    style_formatting = _enrich_style_formatting_from_definitions(doc, style_map, style_formatting) if style_map else style_formatting

    # Build styles: each style name -> full paragraph_format + run_format (font, size, color, alignment, spacing, indent, tab_stops, etc.)
    styles = {}
    for style_name in (para_names or []):
        fmt = style_formatting.get(style_name)
        if not fmt:
            defn = _format_from_style_definition(doc, style_name)
            fmt = defn or {"paragraph_format": {}, "run_format": {}}
        pf = fmt.get("paragraph_format") or {}
        rf = fmt.get("run_format") or {}
        styles[style_name] = {
            "paragraph_format": pf,
            "run_format": rf,
        }

    # Document layout: margins, page size, orientation
    document_layout = _extract_document_layout(doc)

    # Sections: structural roles from template_structure (caption, heading, body, signature_block, etc.)
    template_structure = extract_template_structure(doc, max_paragraphs=500)
    sections = []
    for i, spec in enumerate(template_structure):
        role = spec.get("section_type", "body")
        style_name = spec.get("style", "Normal")
        sections.append({
            "index": i,
            "type": role,
            "style": style_name,
            "block_kind": spec.get("block_kind", "paragraph"),
            "formatting": {
                "paragraph_format": spec.get("paragraph_format") or {},
                "run_format": spec.get("run_format") or {},
            },
            "page_break_before": spec.get("page_break_before", False),
        })

    # Tables: row/col count, style only (no cell text per blueprint rules)
    body = doc.element.body
    qt = qn("w:tbl")
    tables = []
    for table_idx, child in enumerate(body.iterchildren()):
        if child.tag == qt:
            tbl = Table(child, doc)
            rows = len(tbl.rows)
            cols = len(tbl.columns) if tbl.rows else 0
            table_style = None
            try:
                table_style = getattr(tbl, "style", None) and getattr(tbl.style, "name", None)
            except Exception:
                pass
            tables.append({
                "table_index": table_idx,
                "rows": rows,
                "cols": cols,
                "style": table_style,
            })

    # Lists: list-like styles and their formatting (indent, numbering)
    list_styles = []
    for name in (para_names or []):
        if "list" in name.lower() or "number" in name.lower():
            fmt = style_formatting.get(name, {})
            pf = fmt.get("paragraph_format") or {}
            list_styles.append({
                "style_name": name,
                "paragraph_format": pf,
                "run_format": fmt.get("run_format") or {},
            })

    return {
        "document_layout": document_layout,
        "styles": styles,
        "sections": sections,
        "tables": tables,
        "lists": list_styles,
        "style_map": style_map,
        "line_samples": _extract_line_samples(doc),
    }


def extract_styles(doc: Document) -> dict:
    """Extract paragraph style names, style map, per-style formatting, and template content for LLM."""
    para_names = _get_paragraph_style_names(doc)
    if not para_names:
        return {"paragraph_style_names": [], "style_map": {}, "style_formatting": {}}

    heading_like = [
        n for n in para_names
        if any(kw in n.lower() for kw in ("heading", "title", "titre", "section"))
    ]
    list_like = [n for n in para_names if "list" in n.lower() or "number" in n.lower()]

    h1 = _pick_style(para_names, PREFERRED_HEADING_1, heading_like)
    h2 = _pick_style(para_names, PREFERRED_HEADING_2, [n for n in heading_like if n != h1])
    normal = _pick_style(para_names, PREFERRED_NORMAL, para_names)
    list_style = _pick_style(para_names, PREFERRED_LIST, list_like) if list_like else normal

    style_map = {
        "heading": h1 or normal,
        "section_header": h2 or h1 or normal,
        "paragraph": normal,
        "numbered": list_style,
        "wherefore": h2 or h1 or normal,
    }

    style_formatting = _sample_formatting_per_style(doc)
    style_formatting = _enrich_style_formatting_from_definitions(doc, style_map, style_formatting)
    style_formatting = _apply_default_spacing_and_indent(style_map, style_formatting)
    # Prevent template alignment/italic poisoning: body styles get no inherited center or italic
    body_style_names = ("normal", "body text", "list paragraph", "list number", "list")
    for style_name in list(style_formatting.keys()):
        if (style_name or "").lower() in body_style_names:
            fmt = style_formatting[style_name]
            pf = (fmt.get("paragraph_format") or {}).copy()
            pf["alignment"] = None
            rf = (fmt.get("run_format") or {}).copy()
            rf["italic"] = False
            style_formatting[style_name] = {**fmt, "paragraph_format": pf, "run_format": rf}
    line_samples = _extract_line_samples(doc)
    style_guide = build_style_guide(
        style_map=style_map,
        style_formatting=style_formatting,
        all_style_names=para_names,
    )

    template_content = get_template_content_with_styles(doc)
    section_heading_samples = _extract_section_heading_samples(doc)
    template_structure = extract_template_structure(doc)
    tables = extract_tables(doc)

    return {
        "paragraph_style_names": para_names,
        "style_map": style_map,
        "style_formatting": style_formatting,
        "style_guide": style_guide,
        "line_samples": line_samples,
        "template_content": template_content,
        "section_heading_samples": section_heading_samples,
        "template_structure": template_structure,
        "tables": tables,
    }


def save_extracted_styles(schema: dict, base_dir: str = None) -> str:
    """Save extracted style schema to JSON and style guide to plain text. Returns path to JSON file."""
    base_dir = base_dir or os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    store_path = os.path.join(base_dir, STORE_DIR)
    os.makedirs(store_path, exist_ok=True)
    filepath = os.path.join(store_path, EXTRACTED_STYLES_FILE)
    with open(filepath, "w", encoding="utf-8") as f:
        json.dump(schema, f, indent=2, ensure_ascii=False)
    guide = schema.get("style_guide") or schema.get("style_guide_markdown")
    if guide:
        guide_path = os.path.join(store_path, EXTRACTED_STYLE_GUIDE_FILE)
        with open(guide_path, "w", encoding="utf-8") as f:
            f.write(guide)
    return filepath


def load_extracted_styles(base_dir: str = None) -> dict | None:
    """Load extracted style schema from JSON. Returns None if file missing."""
    base_dir = base_dir or os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    filepath = os.path.join(base_dir, STORE_DIR, EXTRACTED_STYLES_FILE)
    if not os.path.isfile(filepath):
        return None
    with open(filepath, encoding="utf-8") as f:
        return json.load(f)


def save_document_blueprint(blueprint: dict, base_dir: str = None) -> str:
    """Save the document blueprint (style/layout schema) to JSON. Returns path to file."""
    base_dir = base_dir or os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    store_path = os.path.join(base_dir, STORE_DIR)
    os.makedirs(store_path, exist_ok=True)
    filepath = os.path.join(store_path, EXTRACTED_BLUEPRINT_FILE)
    with open(filepath, "w", encoding="utf-8") as f:
        json.dump(blueprint, f, indent=2, ensure_ascii=False)
    return filepath


def load_document_blueprint(base_dir: str = None) -> dict | None:
    """Load the document blueprint from JSON. Returns None if file missing."""
    base_dir = base_dir or os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    filepath = os.path.join(base_dir, STORE_DIR, EXTRACTED_BLUEPRINT_FILE)
    if not os.path.isfile(filepath):
        return None
    with open(filepath, encoding="utf-8") as f:
        return json.load(f)
