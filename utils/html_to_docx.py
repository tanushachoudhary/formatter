"""Convert simple HTML (e.g. from Quill editor) to a DOCX document with alignment and inline formatting.
Supports legal-style layout: separator lines (---X), numbered lists with hanging indent, section underlines."""

import re
from html.parser import HTMLParser
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Legal document defaults (match summons / verified complaint style)
DEFAULT_FONT_NAME = "Times New Roman"
DEFAULT_FONT_SIZE_PT = 12.0
LEGAL_SEPARATOR_LINE = "----------------------------------------------------------------------X"

# Quill class ql-font-<value> to display font name (value uses hyphens, e.g. times-new-roman)
_QL_FONT_TO_NAME = {
    "times-new-roman": "Times New Roman",
    "arial": "Arial",
    "calibri": "Calibri",
    "georgia": "Georgia",
    "garamond": "Garamond",
    "verdana": "Verdana",
    "cambria": "Cambria",
}


def _font_from_attrs(attrs) -> str | None:
    """Extract font name from style=font-family or class=ql-font-*."""
    for k, v in attrs:
        if k == "style" and v:
            m = re.search(r"font-family\s*:\s*['\"]?([^;'\"]+)['\"]?", v, re.I)
            if m:
                return m.group(1).strip()
        if k == "class" and v:
            for c in v.split():
                if c.startswith("ql-font-"):
                    val = c[8:].strip()
                    return _QL_FONT_TO_NAME.get(val.lower(), val.replace("-", " ").title())
    return None


def _get_class_set(attrs):
    """Return set of class names from tag attrs."""
    for k, v in attrs:
        if k == "class" and v:
            return set(c.strip() for c in v.split())
    return set()


def _paragraph_border_bottom(paragraph, pt=0.5):
    """Add a thin bottom border to a paragraph (section underline)."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(int(pt * 8)))  # 1/8 pt units
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _looks_like_numbered_paragraph(runs):
    """True if the first run is just digits and a period (e.g. '1. ', '2. ')."""
    if not runs:
        return False
    first_text = (runs[0][0] or "").strip()
    return bool(re.match(r"^\d+\.\s*", first_text))


def _legal_paragraph_format(text: str) -> dict:
    """Infer legal-document formatting (center, bold, italic, underline) from content so download matches summons/complaint style."""
    if not text or not text.strip():
        return {}
    t = text.strip()
    lower = t.lower()
    out = {}
    # Court header: center
    if "supreme court" in lower and ("new york" in lower or "state" in lower) or (t.startswith("COUNTY OF") and len(t) < 50):
        out["alignment"] = "center"
    # Document title: center + bold
    elif t in ("SUMMONS", "VERIFIED COMPLAINT", "COMPLAINT") or (len(t) < 25 and t.isupper() and "cause" not in lower):
        out["alignment"] = "center"
        out["bold"] = True
    # Jury trial / Attorneys for: center + italic
    elif "jury trial demanded" in lower or "attorneys for plaintiff" in lower or "attorneys for defendant" in lower:
        out["alignment"] = "center"
        out["italic"] = True
    # To the above named defendant: center + underline
    elif "to the above named defendant" in lower or "to the above-named defendant" in lower:
        out["alignment"] = "center"
        out["underline"] = True
    # Cause of action / NEGLIGENCE: center + bold
    elif ("as and for" in lower and "cause of action" in lower) or (len(t) < 30 and t.isupper() and not t.endswith(".")):
        out["alignment"] = "center"
        out["bold"] = True
    # -against-: center
    elif t == "-against-" or t.strip() == "-against-":
        out["alignment"] = "center"
    # Firm name / address block (centered): often all caps or has comma/numbers
    elif "pllc" in lower or "p.c." in lower or "esq." in lower:
        if len(t) < 80:
            out["alignment"] = "center"
    return out


def _is_separator_line_only(runs):
    """True if the paragraph is only a legal separator line (dashes/underscores ending in X)."""
    if not runs:
        return False
    text = "".join(r[0] or "" for r in runs).strip()
    if not text or len(text) < 3:
        return False
    if text.endswith("X") or text.endswith("x"):
        text = text[:-1].strip()
    return all(c in " _-.=\t\u00A0" for c in text)


class _SimpleHTMLParser(HTMLParser):
    """Parse HTML into blocks: paragraphs (alignment, runs), list items, separators, underlines."""

    def __init__(self):
        super().__init__()
        self.blocks = []  # each: {block_type, alignment?, runs?, list_item?, list_num?, thin?}
        self._current_runs = []
        self._current_align = None
        self._bold = False
        self._italic = False
        self._underline = False
        self._font_stack = []
        self._in_block = False
        self._current_block_tag = None  # "p", "li", etc.
        self._in_ol = False
        self._list_num = 0

    def _current_font(self):
        return self._font_stack[-1] if self._font_stack else None

    def _start_block(self, attrs, tag="p"):
        self._current_runs = []
        self._current_align = None
        self._current_block_tag = tag
        for k, v in attrs:
            if k == "style" and v:
                m = re.search(r"text-align\s*:\s*(\w+)", v, re.I)
                if m:
                    self._current_align = m.group(1).lower()
        self._in_block = True

    def _end_block(self):
        if not self._in_block:
            return
        if self._current_block_tag == "li":
            self.blocks.append({
                "block_type": "paragraph",
                "alignment": self._current_align,
                "runs": list(self._current_runs),
                "list_item": True,
                "list_num": self._list_num,
            })
        elif self._current_runs:
            self.blocks.append({
                "block_type": "paragraph",
                "alignment": self._current_align,
                "runs": list(self._current_runs),
                "list_item": False,
                "list_num": None,
            })
        self._in_block = False
        self._current_block_tag = None

    def _emit_hr(self, attrs):
        self._end_block()
        classes = _get_class_set(attrs)
        if "thin" in classes or "section-underline" in classes:
            self.blocks.append({"block_type": "underline"})
        else:
            self.blocks.append({"block_type": "separator"})

    def handle_starttag(self, tag, attrs):
        if tag in ("p", "div", "h1", "h2", "h3", "h4", "h5", "h6"):
            self._start_block(attrs, tag=tag)
        elif tag == "hr":
            self._emit_hr(attrs)
        elif tag == "ol":
            self._end_block()
            self._in_ol = True
            self._list_num = 0
        elif tag == "li":
            self._end_block()
            self._list_num += 1
            self._start_block(attrs, tag="li")
        elif tag == "span":
            font = _font_from_attrs(attrs)
            if font:
                self._font_stack.append(font)
        elif tag in ("b", "strong"):
            self._bold = True
        elif tag in ("i", "em"):
            self._italic = True
        elif tag == "u":
            self._underline = True
        elif tag == "br" and self._in_block:
            self._current_runs.append(("\n", self._bold, self._italic, self._underline, self._current_font()))

    def handle_endtag(self, tag):
        if tag in ("p", "div", "h1", "h2", "h3", "h4", "h5", "h6"):
            self._end_block()
        elif tag == "ol":
            self._end_block()
            self._in_ol = False
        elif tag == "li":
            self._end_block()
        elif tag == "span" and self._font_stack:
            self._font_stack.pop()
        elif tag in ("b", "strong"):
            self._bold = False
        elif tag in ("i", "em"):
            self._italic = False
        elif tag == "u":
            self._underline = False

    def handle_data(self, data):
        if self._in_block and data:
            self._current_runs.append((data, self._bold, self._italic, self._underline, self._current_font()))


def html_to_docx_bytes(html: str, font_name: str | None = None, font_size_pt: float | None = None) -> bytes:
    """Convert HTML string to DOCX file bytes. Handles p/div, alignment, b/i/u, <hr>, <ol>/<li>.
    Uses Times New Roman 12pt by default. Legal separators (---X) and section underlines supported."""
    if not (html or "").strip():
        doc = Document()
        doc.add_paragraph()
        out = BytesIO()
        doc.save(out)
        return out.getvalue()

    font_name = font_name or DEFAULT_FONT_NAME
    font_size_pt = font_size_pt if font_size_pt is not None else DEFAULT_FONT_SIZE_PT

    parser = _SimpleHTMLParser()
    html = re.sub(r'class="[^"]*ql-align-(\w+)[^"]*"', r'style="text-align: \1"', html)
    try:
        parser.feed(html)
    except Exception:
        parser.blocks = []

    doc = Document()
    # Legal-style margins: 1 inch all sides (summons / verified complaint)
    try:
        section = doc.sections[-1]
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    except Exception:
        pass

    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    HANG_INDENT = Inches(0.5)
    FIRST_LINE_INDENT = Inches(-0.5)
    for block in parser.blocks:
        block_type = block.get("block_type", "paragraph")

        if block_type == "separator":
            p = doc.add_paragraph()
            run = p.add_run(LEGAL_SEPARATOR_LINE)
            run.font.size = Pt(10)
            run.font.name = font_name
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue

        if block_type == "underline":
            p = doc.add_paragraph()
            _paragraph_border_bottom(p, pt=0.5)
            continue

        # Paragraph (normal or list item)
        runs = block.get("runs") or []
        list_item = block.get("list_item") or False
        if not runs and not list_item:
            continue

        # If paragraph is only a separator line (e.g. pasted "---X"), render as legal separator
        if runs and _is_separator_line_only(runs):
            p = doc.add_paragraph()
            line_text = "".join(r[0] or "" for r in runs).strip()
            run = p.add_run(line_text if line_text else LEGAL_SEPARATOR_LINE)
            run.font.size = Pt(10)
            run.font.name = font_name
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue

        p = doc.add_paragraph()
        full_text = "".join(r[0] or "" for r in runs).strip()
        legal_fmt = _legal_paragraph_format(full_text)
        if block.get("alignment") and block["alignment"] in align_map:
            p.alignment = align_map[block["alignment"]]
        elif legal_fmt.get("alignment") and legal_fmt["alignment"] in align_map:
            p.alignment = align_map[legal_fmt["alignment"]]
        elif not (list_item or _looks_like_numbered_paragraph(runs)) and len(full_text) > 60:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if list_item or _looks_like_numbered_paragraph(runs):
            p.paragraph_format.left_indent = HANG_INDENT
            p.paragraph_format.first_line_indent = FIRST_LINE_INDENT

        for run_tuple in runs:
            text = run_tuple[0]
            bold = run_tuple[1]
            italic = run_tuple[2]
            underline = run_tuple[3]
            inline_font = run_tuple[4] if len(run_tuple) > 4 else None
            if text == "\n":
                p.add_run().add_break()
                continue
            run = p.add_run(text)
            run.bold = bold or legal_fmt.get("bold", False)
            run.italic = italic or legal_fmt.get("italic", False)
            if underline or legal_fmt.get("underline", False):
                run.font.underline = True
            run_font = inline_font or font_name
            if run_font:
                run.font.name = run_font
            run.font.size = Pt(font_size_pt)

    # If the document would open blank (no paragraphs or all empty), add content from raw HTML text
    if not doc.paragraphs or all((p.text or "").strip() == "" for p in doc.paragraphs):
        plain = re.sub(r"<[^>]+>", " ", html).replace("&nbsp;", " ").replace("&amp;", "&").replace("\n", " ").strip()
        plain = " ".join(plain.split())
        if plain:
            doc.add_paragraph(plain).runs[0].font.name = font_name
            doc.paragraphs[-1].runs[0].font.size = Pt(font_size_pt)

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


# Marker for section underlines (solid line under headings) in plain-text preview
SECTION_UNDERLINE_MARKER = "[SECTION_UNDERLINE]"


def plain_text_to_simple_html(text: str) -> str:
    """Wrap plain text in simple HTML paragraphs for the editor.
    Lines that are exactly [SECTION_UNDERLINE] become <hr class="section-underline">."""
    if not text:
        return "<p><br></p>"
    parts = []
    for para in (text or "").split("\n\n"):
        para = (para or "").strip()
        if para == SECTION_UNDERLINE_MARKER:
            parts.append('<hr class="section-underline">')
        else:
            para = (para or "").replace("\n", "<br>")
            if para:
                parts.append(f"<p>{para}</p>")
            else:
                parts.append("<p><br></p>")
    return "".join(parts) if parts else "<p><br></p>"


def simple_html_to_plain_text(html: str) -> str:
    """Extract plain text from simple HTML (for fallback editor).
    <hr class="section-underline"> is emitted as [SECTION_UNDERLINE]."""
    if not html:
        return ""
    # Preserve section underlines as marker
    text = re.sub(
        r'<hr[^>]*class="[^"]*section-underline[^"]*"[^>]*>',
        "\n\n" + SECTION_UNDERLINE_MARKER + "\n\n",
        html,
        flags=re.I,
    )
    # Generic <hr> as double newline
    text = re.sub(r"<hr[^>]*>", "\n\n", text, flags=re.I)
    text = re.sub(r"</p>\s*<p>", "\n\n", text)
    text = text.replace("<p>", "").replace("</p>", "").replace("<br>", "\n")
    return text.strip()
