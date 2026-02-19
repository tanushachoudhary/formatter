import re

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER, WD_UNDERLINE
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# Section underline (thin bottom border) for headings
try:
    from utils.html_to_docx import _paragraph_border_bottom
except Exception:
    _paragraph_border_bottom = None

# Preferred style names to look for in the uploaded document (in order of preference)
PREFERRED_HEADING_1 = ("Heading 1", "Title", "Titre 1")
PREFERRED_HEADING_2 = ("Heading 2", "Subtitle", "Titre 2", "Section")
PREFERRED_NORMAL = ("Normal", "Body Text", "Paragraphe")
PREFERRED_LIST = ("List Number", "List Paragraph", "List")

# Unicode checkbox characters for rendering
CHECKBOX_UNCHECKED = "\u2610"  # ☐
CHECKBOX_CHECKED = "\u2611"    # ☑

# Phrases that indicate address/signature block—do not auto-number these even when using list style
NOT_LIST_CONTENT_PHRASES = (
    "attorneys for plaintiff",
    "attorneys for defendant",
    "park avenue",
    "floor",
    "new york, new york",
    "street",
    "avenue",
    "road",
    "drive",
    "court",
    "pllc",
    "esq.",
    "tel.",
    "fax",
)


# Intro sentences before numbered allegations — keep as body paragraph, never number these
INTRO_PHRASES_NO_NUMBER = (
    "at the time of the accident",
    "at the time of the occurrence",
    "at all times relevant herein",
)


def _looks_like_list_item(text: str) -> bool:
    """True if text looks like a list item (numbered, lettered, or common list starters); False for address/signature/intro."""
    if not text or len(text.strip()) < 3:
        return False
    t = text.strip().lower()
    for phrase in INTRO_PHRASES_NO_NUMBER:
        if t.startswith(phrase):
            return False
    # Only exclude when line starts with address/signature phrases (avoid "court" in "all lower courts" etc.)
    for phrase in NOT_LIST_CONTENT_PHRASES:
        if t.startswith(phrase) and not re.match(r"^[\dai]+[\.\)]\s*", t):
            return False
    if re.match(r"^\(\d{3}\)\s*\d{3}-\d{4}", t):
        return False
    # Numbered or lettered: "1. ...", "a. ...", "i. ..."
    if re.match(r"^\d+[\.\)]\s+", t) or re.match(r"^[a-z][\.\)]\s+", t) or re.match(r"^[ivx]+[\.\)]\s+", t):
        return True
    # Common list starters (any document type)
    list_starts = (
        "that ", "first,", "second,", "third,", "plaintiff ", "plaintiff's ", "defendant ", "the court ",
        "movant ", "respondent ", "applicant ", "petitioner ", "1.", "2.", "a.", "b.",
        "by reason of", "pursuant to", "the detailed", "the above-stated",
    )
    for start in list_starts:
        if t.startswith(start):
            return True
    return False


# Starters for allegation-style paragraphs (so we can split one block into many numbered paragraphs)
ALLEGATION_STARTERS = (
    "that ",
    "by reason of",
    "pursuant to",
    "plaintiff's ",
    "the detailed ",
    "the above-stated ",
)

# NOTICE OF ENTRY / NOTICE OF SETTLEMENT — do not number these even if they start with "that "
NOTICE_ENTRY_SETTLEMENT_STARTERS = (
    "that the within",
    "that an order of which the within",
)


def _is_notice_of_entry_or_settlement(text: str) -> bool:
    """True if paragraph is NOTICE OF ENTRY or NOTICE OF SETTLEMENT text (do not apply list numbering)."""
    if not text or len(text.strip()) < 15:
        return False
    t = text.strip().lower()
    return any(t.startswith(s) for s in NOTICE_ENTRY_SETTLEMENT_STARTERS)


def _starts_allegation(line: str) -> bool:
    """True if line looks like the start of a numbered allegation (e.g. 'That on...', 'By reason of...')."""
    if not line or len(line.strip()) < 10:
        return False
    if _is_notice_of_entry_or_settlement(line):
        return False
    t = line.strip().lower()
    return any(t.startswith(s) for s in ALLEGATION_STARTERS)


def _split_allegation_block(text: str) -> list[str]:
    """If text contains multiple allegation-style paragraphs, split into one string per paragraph for numbering.
    Splits on double newline first; if a chunk contains single newlines and allegation starters, split by line."""
    if not text or not text.strip():
        return []
    text = text.strip()
    # First split by double newline (paragraph boundaries)
    chunks = re.split(r"\n\s*\n", text)
    out = []
    for chunk in chunks:
        chunk = chunk.strip()
        if not chunk:
            continue
        # If this chunk has single newlines and multiple lines that start like allegations, split further
        lines = [ln.strip() for ln in chunk.split("\n") if ln.strip()]
        if len(lines) <= 1:
            out.append(chunk)
            continue
        # Check if we have multiple allegation starters in this chunk
        allegation_lines = [i for i, ln in enumerate(lines) if _starts_allegation(ln)]
        if len(allegation_lines) <= 1:
            out.append(chunk)
            continue
        # Split: each line that starts an allegation begins a new paragraph; merge continuation lines
        current = []
        for ln in lines:
            if _starts_allegation(ln) and current:
                out.append(" ".join(current))
                current = [ln]
            elif _starts_allegation(ln):
                current = [ln]
            else:
                current.append(ln)
        if current:
            out.append(" ".join(current))
    return out if out else [text]


def _get_paragraph_style_names(doc):
    """Return list of paragraph style names defined in the document."""
    return [
        s.name for s in doc.styles
        if s.type == WD_STYLE_TYPE.PARAGRAPH
    ]


def _pick_style(doc, preferred_names, fallback_names=None):
    """Return the first preferred style name that exists in the document, else first fallback."""
    available = _get_paragraph_style_names(doc)
    for name in preferred_names:
        if name in available:
            return name
    if fallback_names:
        for name in fallback_names:
            if name in available:
                return name
    return available[0] if available else None


def _build_style_map_from_doc(doc):
    """Build a block_type -> style_name map using only styles present in the document."""
    para_names = _get_paragraph_style_names(doc)
    if not para_names:
        return None

    # Heading-like: names containing "heading", "title", "titre" (case-insensitive), in stable order
    heading_like = [n for n in para_names if any(kw in n.lower() for kw in ("heading", "title", "titre", "section"))]
    # List-like
    list_like = [n for n in para_names if "list" in n.lower() or "number" in n.lower()]

    h1 = _pick_style(doc, PREFERRED_HEADING_1, heading_like)
    h2 = _pick_style(doc, PREFERRED_HEADING_2, [n for n in heading_like if n != h1])
    normal = _pick_style(doc, PREFERRED_NORMAL, para_names)
    list_style = _pick_style(doc, PREFERRED_LIST, list_like) if list_like else normal

    return {
        "heading": h1 or normal,
        "section_header": h2 or h1 or normal,
        "paragraph": normal,
        "numbered": list_style,
        "wherefore": h2 or h1 or normal,
    }


# Style names that are body text — always justify, never center; never inherit template italic
BODY_STYLE_NAMES = ("normal", "body text", "list paragraph", "list number", "list")

def _block_type_for_alignment(block_kind: str, section_type: str, style_name: str = "") -> str:
    """Map block_kind + section_type to alignment block_type for enforce_legal_alignment."""
    if block_kind == "line":
        return "line"
    if block_kind == "signature_line":
        return "signature"
    if block_kind == "section_underline":
        return "paragraph"
    # Body styles: always justify (prevent template center/italic from leaking)
    if (style_name or "").strip().lower() in BODY_STYLE_NAMES:
        return "paragraph"
    # Content slots: use section_type
    if section_type in ("caption",):
        return "section_header"
    if section_type in ("attorney_signature", "notary"):
        return "signature"
    if section_type == "to_section":
        return "to_section"
    return "paragraph"


def enforce_legal_alignment(block_type: str, paragraph):
    """Override alignment: left for captions/headings/signatures; justify for body; consistent legal layout."""
    if not paragraph:
        return
    try:
        if block_type in ("heading", "section_header"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif block_type in ("paragraph", "numbered", "body"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif block_type in ("signature", "address", "to_section"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif block_type == "line":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    except Exception:
        pass


def clear_body_italic(paragraph):
    """Remove italic from all runs so body text is not italicised by template style."""
    if not paragraph:
        return
    try:
        for run in paragraph.runs:
            run.italic = False
    except Exception:
        pass


def force_legal_run_format(paragraph):
    """Force black color and no italic on all runs (avoids blue/hyperlink/italic from template styles)."""
    if not paragraph:
        return
    try:
        for run in paragraph.runs:
            try:
                run.font.color.rgb = RGBColor(0, 0, 0)
            except Exception:
                pass
            try:
                run.font.italic = False
            except Exception:
                pass
    except Exception:
        pass


def force_legal_run_format_document(doc):
    """Force black color and no italic on every paragraph in the document."""
    if not doc:
        return
    try:
        for paragraph in doc.paragraphs:
            force_legal_run_format(paragraph)
    except Exception:
        pass


def _apply_paragraph_format(paragraph, fmt: dict):
    """Apply stored paragraph format dict (exact Word features: alignment, spacing, indent, line_spacing, keep_*, page_break_before)."""
    if not fmt or not paragraph:
        return
    pf = paragraph.paragraph_format
    try:
        if "alignment" in fmt and fmt["alignment"]:
            alignment = getattr(WD_ALIGN_PARAGRAPH, fmt["alignment"], None)
            if alignment is not None:
                pf.alignment = alignment
    except Exception:
        pass
    for attr, key in (
        ("space_before", "space_before"),
        ("space_after", "space_after"),
        ("left_indent", "left_indent"),
        ("right_indent", "right_indent"),
        ("first_line_indent", "first_line_indent"),
    ):
        try:
            val = fmt.get(key)
            if val is not None and isinstance(val, (int, float)):
                setattr(pf, attr, Pt(val))
        except Exception:
            pass
    try:
        if "line_spacing" in fmt and fmt["line_spacing"] is not None:
            val = fmt["line_spacing"]
            rule_name = fmt.get("line_spacing_rule")
            rule = getattr(WD_LINE_SPACING, rule_name, None) if isinstance(rule_name, str) else None
            # EXACTLY or AT_LEAST: use fixed height in points
            if rule in (WD_LINE_SPACING.EXACTLY, WD_LINE_SPACING.AT_LEAST):
                pf.line_spacing = Pt(val) if isinstance(val, (int, float)) else val
                pf.line_spacing_rule = rule
            # MULTIPLE, SINGLE, DOUBLE, ONE_POINT_FIVE: use multiplier (float)
            else:
                num = float(val) if isinstance(val, (int, float)) else None
                if num is not None and 0.25 <= num <= 3.0:
                    pf.line_spacing = num
    except Exception:
        pass
    for attr in ("page_break_before", "keep_with_next", "keep_together"):
        try:
            if attr in fmt and fmt[attr] is not None:
                setattr(pf, attr, bool(fmt[attr]))
        except Exception:
            pass
    try:
        tab_stops = fmt.get("tab_stops")
        if tab_stops and isinstance(tab_stops, list):
            pf.tab_stops.clear_all()
            for ts in tab_stops:
                pos_pt = ts.get("position_pt") if isinstance(ts, dict) else None
                if pos_pt is None:
                    continue
                align_name = (ts.get("alignment") or "LEFT") if isinstance(ts, dict) else "LEFT"
                leader_name = (ts.get("leader") or "SPACES") if isinstance(ts, dict) else "SPACES"
                align = getattr(WD_TAB_ALIGNMENT, align_name, WD_TAB_ALIGNMENT.LEFT)
                leader = getattr(WD_TAB_LEADER, leader_name, WD_TAB_LEADER.SPACES)
                pf.tab_stops.add_tab_stop(Pt(pos_pt), align, leader)
    except Exception:
        pass


def _render_checkboxes(text: str) -> str:
    """Replace [ ], [x], [X] with Unicode checkbox characters so they render in the document."""
    if not text:
        return text
    text = re.sub(r"\[\s*[xX]\s*\]", CHECKBOX_CHECKED + " ", text)
    text = re.sub(r"\[\s*\]", CHECKBOX_UNCHECKED + " ", text)
    return text


def _is_section_start(
    text: str, block_type: str, style_map: dict, valid_style_names: set,
    section_heading_samples: list = None,
) -> bool:
    """True if this heading should get a page break (template-driven: only when template had page break before this text)."""
    if not text or not text.strip():
        return False
    is_heading = (
        block_type in ("heading", "section_header")
        or (style_map.get("heading") and block_type == style_map["heading"])
        or (style_map.get("section_header") and block_type == style_map["section_header"])
    )
    if not is_heading:
        return False
    if not section_heading_samples:
        return False
    t = text.strip().lower()
    for sample in section_heading_samples:
        if sample in t or t in sample or t.startswith(sample) or sample.startswith(t):
            return True
    return False


def _add_paragraph_with_inline_formatting(doc, segments: list[tuple[str, bool, bool]], style, run_fmt_base: dict):
    """Add a paragraph with multiple runs for bold/italic segments. style and run_fmt_base from template."""
    p = doc.add_paragraph(style=style)
    for seg_text, bold, italic in segments:
        if not seg_text:
            continue
        run = p.add_run(seg_text)
        fmt = dict(run_fmt_base)
        if bold:
            fmt["bold"] = True
        if italic:
            fmt["italic"] = True
        _apply_run_format(run, fmt)
    return p


def _apply_num_pr(paragraph, num_id: int, ilvl: int = 0):
    """Set Word list numbering on a paragraph (numPr) so it displays as 1., 2., 3."""
    if not paragraph or num_id is None:
        return
    try:
        p_el = paragraph._element
        pPr = p_el.get_or_add_pPr()
        numPr = OxmlElement("w:numPr")
        numId_el = OxmlElement("w:numId")
        numId_el.set(qn("w:val"), str(num_id))
        numPr.append(numId_el)
        ilvl_el = OxmlElement("w:ilvl")
        ilvl_el.set(qn("w:val"), str(ilvl))
        numPr.append(ilvl_el)
        pPr.append(numPr)
    except Exception:
        pass


def _apply_run_format(run, fmt: dict):
    """Apply stored run/font format (bold, italic, underline, font name/size). Color is not applied so output stays black (legal standard)."""
    if not fmt or not run:
        return
    font = run.font
    try:
        if "bold" in fmt:
            font.bold = fmt["bold"]
    except Exception:
        pass
    try:
        if "italic" in fmt:
            font.italic = fmt["italic"]
    except Exception:
        pass
    try:
        if "underline" in fmt:
            u = fmt["underline"]
            if u is True or u == "True":
                font.underline = True
            elif u is False or u == "False":
                font.underline = False
            elif isinstance(u, str) and hasattr(WD_UNDERLINE, u):
                font.underline = getattr(WD_UNDERLINE, u)
            else:
                font.underline = u
    except Exception:
        pass
    try:
        if "name" in fmt and fmt["name"]:
            font.name = fmt["name"]
    except Exception:
        pass
    try:
        if "size_pt" in fmt and fmt["size_pt"] is not None:
            font.size = Pt(fmt["size_pt"])
    except Exception:
        pass
    # Force black text and no italic (legal standard); do not copy template color (e.g. blue) or italic
    try:
        font.color.rgb = RGBColor(0, 0, 0)
    except Exception:
        pass
    try:
        font.italic = False
    except Exception:
        pass


# Fallback when template has no line samples
DEFAULT_SIGNATURE_LINE = "_________________________"
# Default separator line (dashes ending in X) so it always renders
DEFAULT_LINE = "----------------------------------------------------------------------X"


def _is_separator_noise(text: str) -> bool:
    """True if text is only underscores, dashes, equals, spaces, dots, or ends with X (stray separator noise)."""
    if not text or not text.strip():
        return True
    t = text.strip()
    if not t:
        return True
    # Allow trailing X (legal separator style e.g. "------------------------------------------------------------------X")
    if t and t[-1] in ("X", "x"):
        t = t[:-1].strip()
    # Only these chars: space, underscore, hyphen, dot, equals
    allowed = set(" _-.=\u00A0\t")
    if all(c in allowed for c in t):
        return True
    if re.match(r"^[\s\-\._=]+$", t):
        return True
    return False

# Phrases that start the main body (after caption); caption = everything before this
BODY_START_PHRASES = (
    "please take notice",
    "take further notice",
    "dated:",
    "affirms the following",
    "under the penalties of perjury",
    "being duly sworn",
    "duly sworn, says",
)
# Right-column caption: index number and motion/document title
RIGHT_CAPTION_PHRASES = (
    "index no",
    "index number",
    "notice of motion",
    "to restore",
    "affirmation in support",
    "affidavit of service",
    "memorandum of law",
)
# Block text that starts a new document (repeated caption) when not at start of paste
NEW_DOCUMENT_START_PHRASES = (
    "supreme court of the state of new york",
    "supreme court of new york",
)

# Court caption patterns: use one consistent style for all so layout doesn't vary
COURT_CAPTION_PHRASES = (
    "superior court of",
    "supreme court of",
    "index no.:",
    "index no.",
    "date filed:",
    "plaintiff,",
    "defendant.",
    "-against-",
)

# Phrases that start a major section: add space_before for clear separation
SECTION_STARTER_PHRASES = (
    "to the above named defendant",
    "wherefore",
    "dated",
    "yours, etc",
    "please take notice",
)
# Space before a section starter (pt) and after court caption (pt) for consistent structure
SPACE_BEFORE_SECTION_PT = 12.0
SPACE_AFTER_CAPTION_PT = 6.0

# Numbered list (allegations): spacing and hanging indent for clean legal layout
SPACE_BEFORE_NUMBERED_PT = 0.0
SPACE_AFTER_NUMBERED_PT = 8.0   # space between each numbered point for readability
NUMBERED_LEFT_INDENT_PT = 18.0   # body text indented 0.25"
NUMBERED_FIRST_LINE_INDENT_PT = -18.0  # hanging: number left-aligned, description indented

# Cause-of-action headings (e.g. "AS AND FOR A FIRST CAUSE OF ACTION:") — treat as section header
CAUSE_OF_ACTION_PHRASE = "cause of action"


def _looks_like_court_caption(text: str) -> bool:
    """True if block text is a court caption line (so we can apply one consistent style)."""
    if not text or len(text.strip()) < 3:
        return False
    t = text.strip().lower()
    return any(p in t or t.startswith(p) for p in COURT_CAPTION_PHRASES)


def _is_section_starter(text: str) -> bool:
    """True if paragraph starts a major section (TO THE ABOVE NAMED DEFENDANT, WHEREFORE, Dated, etc.)."""
    if not text or len(text.strip()) < 4:
        return False
    t = text.strip().lower()
    return any(t.startswith(p) or t == p for p in SECTION_STARTER_PHRASES)


def _looks_like_cause_of_action_heading(text: str) -> bool:
    """True if paragraph is a cause-of-action heading (e.g. 'AS AND FOR A FIRST CAUSE OF ACTION:')."""
    if not text or len(text.strip()) < 10:
        return False
    t = text.strip().lower()
    return CAUSE_OF_ACTION_PHRASE in t and "as and for" in t


def _apply_numbered_paragraph_layout(paragraph):
    """Apply consistent spacing and hanging indent for numbered points: number left-aligned, description indented."""
    if not paragraph:
        return
    try:
        pf = paragraph.paragraph_format
        pf.space_before = Pt(SPACE_BEFORE_NUMBERED_PT)
        pf.space_after = Pt(SPACE_AFTER_NUMBERED_PT)
        pf.left_indent = Pt(NUMBERED_LEFT_INDENT_PT)
        pf.first_line_indent = Pt(NUMBERED_FIRST_LINE_INDENT_PT)
    except Exception:
        pass


def _apply_section_spacing(paragraph, text: str, is_court_caption: bool):
    """Set space_before for section starters and space_after for caption lines so sections are well-separated."""
    if not paragraph:
        return
    try:
        pf = paragraph.paragraph_format
        if _is_section_starter(text):
            pf.space_before = Pt(SPACE_BEFORE_SECTION_PT)
        if _looks_like_cause_of_action_heading(text):
            pf.space_before = Pt(SPACE_BEFORE_SECTION_PT)
            pf.space_after = Pt(SPACE_AFTER_CAPTION_PT)
        if is_court_caption:
            pf.space_after = Pt(SPACE_AFTER_CAPTION_PT)
    except Exception:
        pass


def _split_into_document_segments(blocks: list) -> list[list]:
    """Split blocks into one list per document. New document starts at repeated court heading (not at index 0)."""
    if not blocks:
        return []
    segment_starts = [0]
    for i in range(1, len(blocks)):
        bt, text = blocks[i]
        t = (text or "").strip().lower()
        if not t:
            continue
        for phrase in NEW_DOCUMENT_START_PHRASES:
            if t.startswith(phrase) or t == phrase.strip() or phrase in t[:80]:
                segment_starts.append(i)
                break
    out = []
    for j in range(len(segment_starts)):
        start = segment_starts[j]
        end = segment_starts[j + 1] if j + 1 < len(segment_starts) else len(blocks)
        out.append(blocks[start:end])
    return out


def _split_caption_body(blocks: list) -> tuple[list, list, list]:
    """Split blocks into caption_left, caption_right, body. Returns (caption_left, caption_right, body_blocks)."""
    if not blocks:
        return [], [], []
    body_start_idx = None
    for i, (bt, text) in enumerate(blocks):
        t = (text or "").strip().lower()
        if not t:
            continue
        for phrase in BODY_START_PHRASES:
            if phrase in t or t.startswith(phrase):
                body_start_idx = i
                break
        if body_start_idx is not None:
            break
    if body_start_idx is None:
        return [], [], blocks
    caption_blocks = blocks[:body_start_idx]
    body_blocks = blocks[body_start_idx:]
    left, right = [], []
    for b in caption_blocks:
        bt, text = b
        t = (text or "").strip().lower()
        is_right = any(p in t for p in RIGHT_CAPTION_PHRASES) or (t == "to restore")
        if is_right:
            right.append(b)
        else:
            left.append(b)
    return left, right, body_blocks


def _resolve_style(block_type: str, style_map: dict, style_formatting: dict):
    """Resolve block_type to a style name: use template style name if present, else logical style_map."""
    if block_type in style_formatting:
        return block_type
    return style_map.get(block_type, style_map.get("paragraph"))


def inject_blocks(doc, blocks, style_map=None, style_formatting=None, line_samples=None, section_heading_samples=None, template_structure=None, numbered_num_id=None, numbered_ilvl=0):
    """Inject text into template structure. When template_structure is provided (slot-fill):
    assign paragraph.style = template style only — no manual formatting. Word handles layout,
    numbering, spacing from style definitions. Renderer never invents formatting.
    numbered_num_id/numbered_ilvl: when set, paragraphs with the numbered style get list numbering (1., 2., 3.)."""
    if style_map is None:
        style_map = _build_style_map_from_doc(doc)
    if not style_map:
        style_map = {"heading": None, "section_header": None, "paragraph": None, "numbered": None, "wherefore": None}
    style_formatting = style_formatting or {}
    line_samples = line_samples or []
    section_heading_samples = section_heading_samples or []
    valid_style_names = set(style_formatting.keys())

    # Structure-driven slot-fill: parser only — assign existing text to slots; never invent or fallback.
    if template_structure and len(blocks) == len(template_structure):
        seen = set()  # Caption deduplication: do not render the same block text twice (stops repeated court headers)
        for i in range(len(template_structure)):
            spec = template_structure[i]
            style = (blocks[i][0] if isinstance(blocks[i], (list, tuple)) else spec.get("style", "Normal"))
            slot_text = (blocks[i][1] if isinstance(blocks[i], (list, tuple)) and len(blocks[i]) > 1 else (blocks[i] if isinstance(blocks[i], str) else ""))
            slot_text = (slot_text or "").strip()
            if style not in valid_style_names:
                style = style_map.get("paragraph") or (list(valid_style_names)[0] if valid_style_names else "Normal")
            block_kind = spec.get("block_kind", "paragraph")
            section_type = spec.get("section_type", "body")
            template_text = (spec.get("template_text") or "").strip()
            if spec.get("page_break_before"):
                doc.add_page_break()
            if block_kind == "line":
                if not (slot_text or template_text).strip():
                    continue
                if template_text:
                    p = doc.add_paragraph(template_text, style=style)
                    fmt = (style_formatting.get(style) or {}).get("paragraph_format") or {}
                    _apply_paragraph_format(p, fmt)
                    enforce_legal_alignment(_block_type_for_alignment(block_kind, section_type, style), p)
                continue
            if block_kind == "section_underline":
                p = doc.add_paragraph(style=style)
                if _paragraph_border_bottom:
                    _paragraph_border_bottom(p, pt=0.5)
                fmt = (style_formatting.get(style) or {}).get("paragraph_format") or {}
                _apply_paragraph_format(p, fmt)
                enforce_legal_alignment(_block_type_for_alignment(block_kind, section_type, style), p)
                continue
            if block_kind == "signature_line":
                if template_text:
                    p = doc.add_paragraph(template_text, style=style)
                    fmt = (style_formatting.get(style) or {}).get("paragraph_format") or {}
                    _apply_paragraph_format(p, fmt)
                    enforce_legal_alignment(_block_type_for_alignment(block_kind, section_type, style), p)
                continue
            # Content slots: empty → skip; dedupe then render
            if not slot_text:
                continue
            if slot_text in seen:
                continue
            seen.add(slot_text)
            p = doc.add_paragraph(style=style)
            p.add_run(_render_checkboxes(slot_text))
            fmt = (style_formatting.get(style) or {}).get("paragraph_format") or {}
            _apply_paragraph_format(p, fmt)
            align_type = _block_type_for_alignment(block_kind, section_type, style)
            enforce_legal_alignment(align_type, p)
            if align_type == "paragraph":
                clear_body_italic(p)
        trim_trailing_separators(doc)
        return

    # Fallback path when no template_structure: still use style only; no fake numbering (Word handles via style).
    # Deduplicate long repeated blocks (e.g. same summons/caption pasted multiple times) so output isn't bloated.
    MIN_DEDUP_LEN = 80  # Only skip when this many chars and we've seen this exact text before
    seen_long_text = set()

    segments = _split_into_document_segments(blocks)
    for seg_idx, segment in enumerate(segments):
        if seg_idx > 0:
            doc.add_page_break()
        caption_left, caption_right, body_blocks = _split_caption_body(segment)
        blocks_to_render = caption_left + caption_right + body_blocks if (caption_left or caption_right) else segment
        section_break_added_in_segment = False

        for block_type, text in blocks_to_render:
            text = (text or "").strip()

            # Skip long duplicate paragraphs (repeated summons, captions, allegations from concatenated input)
            if len(text) >= MIN_DEDUP_LEN:
                normalized = re.sub(r"\s+", " ", text).strip()
                if normalized in seen_long_text:
                    continue
                seen_long_text.add(normalized)

            if block_type == "page_break":
                doc.add_page_break()
                continue

            if block_type == "signature_line":
                label = (text.strip() if text and text.strip() and text.strip() not in ("---", "—", "-") else None)
                line_text = None
                if line_samples:
                    for s in line_samples:
                        t = s.get("text", "")
                        if "_" in t and t.strip().replace("_", "").replace(" ", "") == "":
                            line_text = t
                            break
                    if line_text is None:
                        line_text = line_samples[0].get("text", DEFAULT_SIGNATURE_LINE)
                if not line_text:
                    line_text = DEFAULT_SIGNATURE_LINE
                if label:
                    line_text = f"{line_text}  {label}"
                style = _resolve_style("paragraph", style_map, style_formatting)
                p = doc.add_paragraph(line_text, style=style)
                fmt = (style_formatting.get(style) or {}).get("paragraph_format") or {}
                _apply_paragraph_format(p, fmt)
                enforce_legal_alignment("signature", p)
                continue

            if block_type == "section_underline":
                style = _resolve_style("paragraph", style_map, style_formatting)
                p = doc.add_paragraph(style=style)
                if _paragraph_border_bottom:
                    _paragraph_border_bottom(p, pt=0.5)
                fmt = (style_formatting.get(style) or {}).get("paragraph_format") or {}
                _apply_paragraph_format(p, fmt)
                enforce_legal_alignment("paragraph", p)
                continue

            if block_type == "line":
                line_text = (text or "").strip()
                if line_text and ("block_type" in line_text or "text field" in line_text):
                    line_text = ""
                if not line_text and line_samples:
                    for s in line_samples:
                        t = s.get("text", "")
                        if t.rstrip().endswith("X") or t.rstrip().endswith("x"):
                            line_text = t
                            break
                    if not line_text:
                        line_text = line_samples[0].get("text", DEFAULT_LINE)
                if not line_text:
                    line_text = DEFAULT_LINE
                style = _resolve_style("paragraph", style_map, style_formatting)
                p = doc.add_paragraph(line_text, style=style)
                fmt = (style_formatting.get(style) or {}).get("paragraph_format") or {}
                _apply_paragraph_format(p, fmt)
                enforce_legal_alignment("line", p)
                continue

            if not text:
                continue

            # Split one block into multiple numbered paragraphs when it contains many allegations (e.g. paste of "That on...", "By reason of...")
            numbered_style = style_map.get("numbered") and (not valid_style_names or style_map["numbered"] in valid_style_names)
            first_line = text.split("\n")[0].strip() if "\n" in text else text
            lines_in_block = [ln.strip() for ln in text.split("\n") if ln.strip()]
            has_any_allegation = any(_starts_allegation(ln) for ln in lines_in_block)
            allegation_paras = _split_allegation_block(text) if numbered_style and (_looks_like_list_item(first_line) or (len(lines_in_block) > 1 and has_any_allegation)) else []

            if len(allegation_paras) > 1:
                # Render each allegation as its own numbered paragraph. Do NOT hardcode "1.", "2." as text:
                # strip any leading number from content and apply Word numPr so the template controls numbering.
                style = style_map["numbered"]
                for one in allegation_paras:
                    one = one.strip()
                    if not one:
                        continue
                    one = re.sub(r"^\d+[\.\)]\s*", "", one).strip()
                    one = re.sub(r"^[a-z][\.\)]\s*", "", one, count=1).strip()
                    one = re.sub(r"^[ivx]+[\.\)]\s*", "", one, count=1, flags=re.IGNORECASE).strip()
                    one = _render_checkboxes(one)
                    segments = [(one, False, False)]
                    _add_paragraph_with_inline_formatting(doc, segments, style, {})
                    p = doc.paragraphs[-1] if doc.paragraphs else None
                    if p:
                        fmt = (style_formatting.get(style) or {}).get("paragraph_format") or {}
                        _apply_paragraph_format(p, fmt)
                        # Number only negligence-style allegations (That on..., By reason of..., The above-stated...), not WHEREFORE demands or other points
                        if numbered_num_id is not None and _starts_allegation(one):
                            _apply_num_pr(p, numbered_num_id, numbered_ilvl)
                        if _starts_allegation(one):
                            _apply_numbered_paragraph_layout(p)
                        enforce_legal_alignment("numbered", p)
                        clear_body_italic(p)
                continue

            is_court_caption = _looks_like_court_caption(text)
            is_cause_of_action_heading = _looks_like_cause_of_action_heading(text)
            # Use one consistent style for court caption lines; cause-of-action headings get section_header for clear distinction from numbered points
            if is_court_caption:
                style = style_map.get("section_header") or style_map.get("heading") or style_map.get("paragraph")
                if not style or (valid_style_names and style not in valid_style_names):
                    style = list(valid_style_names)[0] if valid_style_names else "Normal"
            elif is_cause_of_action_heading:
                style = style_map.get("section_header") or style_map.get("heading") or style_map.get("paragraph")
                if not style or (valid_style_names and style not in valid_style_names):
                    style = list(valid_style_names)[0] if valid_style_names else "Normal"
            elif _looks_like_list_item(text) and numbered_style:
                # Use numbered/list style so Word numbers allegations (1., 2., 3.)
                style = style_map["numbered"]
            else:
                style = block_type if block_type in valid_style_names else style_map.get(block_type, style_map.get("paragraph"))
            # If LLM gave paragraph/body but content is an allegation (That on..., By reason of...), use numbered style so it gets numbered
            if style != style_map.get("numbered") and _starts_allegation((text or "").strip()) and numbered_style:
                style = style_map["numbered"]
            if not style:
                style = list(valid_style_names)[0] if valid_style_names else "Normal"
            # Only one page break per segment for "section start"
            if doc.paragraphs and not section_break_added_in_segment and _is_section_start(text, block_type, style_map, valid_style_names, section_heading_samples):
                doc.add_page_break()
                section_break_added_in_segment = True
            # Do NOT hardcode numbers: strip leading "1." etc. so Word (via numPr/template style) supplies the number
            if _looks_like_list_item(text):
                text = re.sub(r"^\d+[\.\)]\s*", "", text).strip()
                text = re.sub(r"^[a-z][\.\)]\s*", "", text, count=1).strip()
                text = re.sub(r"^[ivx]+[\.\)]\s*", "", text, count=1, flags=re.IGNORECASE).strip()
            text = _render_checkboxes(text)
            segments = [(text, False, False)]
            _add_paragraph_with_inline_formatting(doc, segments, style, {})
            p = doc.paragraphs[-1] if doc.paragraphs else None
            if p:
                fmt = (style_formatting.get(style) or {}).get("paragraph_format") or {}
                _apply_paragraph_format(p, fmt)
                # Number only negligence-style allegations (That on..., By reason of..., The above-stated...), not WHEREFORE or other points
                txt_stripped = (text or "").strip()
                is_negligence_allegation = style == style_map.get("numbered") and _starts_allegation(txt_stripped)
                if is_negligence_allegation and numbered_num_id is not None:
                    _apply_num_pr(p, numbered_num_id, numbered_ilvl)
                _apply_section_spacing(p, txt_stripped, is_court_caption=is_court_caption)
                if is_negligence_allegation:
                    _apply_numbered_paragraph_layout(p)
                align_type = "section_header" if style in (style_map.get("heading"), style_map.get("section_header")) else ("numbered" if is_negligence_allegation else "paragraph")
                enforce_legal_alignment(align_type, p)
                if align_type == "paragraph" or align_type == "numbered":
                    clear_body_italic(p)
    trim_trailing_separators(doc)


def _is_empty_or_noise_paragraph(para) -> bool:
    """True if paragraph has no content or only separator noise (whitespace, underscores, '- - -')."""
    if not para:
        return True
    text = (para.text or "").strip()
    if not text:
        return True
    return _is_separator_noise(text)


def trim_trailing_separators(doc):
    """Remove trailing paragraphs that look like separators (----, ====, ______). Call after rendering, before save."""
    def is_separator(text):
        t = (text or "").strip()
        return t.startswith("-") or t.startswith("=") or t.startswith("_")

    while doc.paragraphs:
        last = doc.paragraphs[-1]
        if is_separator(last.text):
            try:
                p = last._element
                p.getparent().remove(p)
            except Exception:
                break
        else:
            break


def remove_trailing_empty_and_noise(doc):
    """Remove trailing paragraphs that are empty or only separator noise (underscores, '- - -')."""
    paras = list(doc.paragraphs)
    if not paras:
        return
    removed = 0
    for para in reversed(paras):
        if _is_empty_or_noise_paragraph(para):
            try:
                p_el = para._element
                p_el.getparent().remove(p_el)
                removed += 1
            except Exception:
                break
        else:
            break


def force_single_column(doc):
    """Force all sections to single-column layout so the document renders as one column per page, not multi-column.
    Handles sectPr as direct children of body and sectPr inside paragraph properties (section breaks)."""
    try:
        body = doc.element.body
        for sect_pr in body.iter(qn("w:sectPr")):
            cols = None
            for c in sect_pr:
                if c.tag == qn("w:cols"):
                    cols = c
                    break
            if cols is not None:
                cols.set(qn("w:num"), "1")
            else:
                cols_el = OxmlElement("w:cols")
                cols_el.set(qn("w:num"), "1")
                sect_pr.insert(0, cols_el)
    except Exception:
        pass


def clear_document_body(doc):
    """Remove all paragraphs and tables from the document body, keeping section properties."""
    for para in list(doc.paragraphs):
        p_el = para._element
        p_el.getparent().remove(p_el)
    for table in list(doc.tables):
        tbl_el = table._element
        tbl_el.getparent().remove(tbl_el)
