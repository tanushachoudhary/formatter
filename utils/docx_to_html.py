"""Convert DOCX to HTML for web rendering or round-trip editing.

Uses Mammoth when available to produce clean, semantic HTML that preserves
structure and basic formatting. Falls back to plain-text extraction if Mammoth
is not installed.
"""

from __future__ import annotations

from io import BytesIO
from typing import BinaryIO


def docx_to_html(docx_input: str | bytes | BinaryIO) -> str:
    """
    Convert a DOCX document to HTML.

    Args:
        docx_input: Path to a .docx file, or bytes, or a file-like object opened in binary mode.

    Returns:
        HTML string (e.g. from Mammoth) or a simple HTML wrapper around extracted text if Mammoth is missing.
    """
    try:
        import mammoth
    except ImportError:
        return _docx_to_html_fallback(docx_input)

    if isinstance(docx_input, str):
        with open(docx_input, "rb") as f:
            result = mammoth.convert_to_html(f)
    elif isinstance(docx_input, bytes):
        result = mammoth.convert_to_html(BytesIO(docx_input))
    else:
        result = mammoth.convert_to_html(docx_input)

    return result.value or ""


def _docx_to_html_fallback(docx_input: str | bytes | BinaryIO) -> str:
    """Fallback: extract paragraphs as <p> when Mammoth is not installed."""
    from docx import Document

    if isinstance(docx_input, str):
        doc = Document(docx_input)
    elif isinstance(docx_input, bytes):
        doc = Document(BytesIO(docx_input))
    else:
        doc = Document(docx_input)

    parts = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            parts.append("<p><br></p>")
        else:
            text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            parts.append(f"<p>{text}</p>")
    return "\n".join(parts) if parts else "<p><br></p>"
