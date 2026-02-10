import streamlit as st
from docx import Document
import shutil
import os

st.title("DOCX Renderer Debugger")

# ---- CONFIG ----
TEMPLATE_PATH = "pleading_template.docx"
OUTPUT_PATH = "generated.docx"


import re

def detect_style_roles(doc):

    roles = {
        "caption": None,
        "allegation": None,
        "heading": None,
        "body": None
    }

    for p in doc.paragraphs:

        text = p.text.strip()

        if "SUPREME COURT" in text:
            roles["caption"] = p.style.name

        elif re.match(r"^\d+\.", text):
            roles["allegation"] = p.style.name

        elif text.isupper():
            roles["heading"] = p.style.name

        elif roles["body"] is None:
            roles["body"] = p.style.name

    return roles


# ---- UTILITIES ----
def extract_template_styles(doc):
    styles = {}

    for s in doc.styles:
        styles[s.name] = {
            "font": s.font.name,
            "size": s.font.size,
            "bold": s.font.bold
        }

    return styles

def render_document(doc, raw_text, roles):

    blocks = parse_text(raw_text)

    allegations = []

    for kind, text in blocks:

        if kind == "heading":
            doc.add_paragraph(text, style=roles["heading"])

        elif kind == "allegation":
            allegations.append(text)

        else:
            doc.add_paragraph(text, style=roles["body"])

    for a in allegations:
        doc.add_paragraph(a, style=roles["allegation"])


def clone_template():
    if not os.path.exists(TEMPLATE_PATH):
        st.error("❌ Template file not found!")
        return None

    shutil.copy(TEMPLATE_PATH, OUTPUT_PATH)
    return Document(OUTPUT_PATH)


def debug_styles(doc):
    st.subheader("📋 Styles Found in Template")

    style_names = []
    for s in doc.styles:
        style_names.append(s.name)

    st.write(style_names)


def generate_document():
    doc = clone_template()
    if not doc:
        return

    # DEBUG — list styles
    debug_styles(doc)

    # Try applying styles
    try:
        doc.add_paragraph(
            "SUPREME COURT OF THE STATE OF NEW YORK",
            style="Normal"   # Change later to your pleading style
        )

        doc.add_paragraph(
            "Debug paragraph with template styles.",
            style="Normal"
        )

        doc.save(OUTPUT_PATH)
        st.success("✅ Document generated successfully!")

    except Exception as e:
        st.error(f"❌ Style application failed: {e}")


# ---- UI ----

if st.button("Generate DOCX"):
    generate_document()

    if os.path.exists(OUTPUT_PATH):
        with open(OUTPUT_PATH, "rb") as f:
            st.download_button(
                label="⬇ Download Generated DOCX",
                data=f,
                file_name="generated.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
